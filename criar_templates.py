from docx import Document
from docx.shared import Pt
import os

def criar_template_contrato_servicos():
    doc = Document()
    
    # Título
    titulo = doc.add_heading('CONTRATO DE PRESTAÇÃO DE SERVIÇOS', 0)
    titulo.alignment = 1  # Centralizado
    
    # Parágrafos
    doc.add_paragraph('Pelo presente instrumento particular de contrato de prestação de serviços, de um lado,')
    
    # Campos para preenchimento
    doc.add_paragraph('CONTRATANTE: {{ nome_cliente }}')
    doc.add_paragraph('CPF: {{ cpf }}')
    doc.add_paragraph('CONTRATADO: [NOME DA EMPRESA]')
    doc.add_paragraph('VALOR DO CONTRATO: R$ {{ valor_contrato }}')
    
    # Cláusulas
    doc.add_heading('CLÁUSULAS DO CONTRATO', level=1)
    
    clausulas = [
        "1. DO OBJETO: O presente contrato tem por objeto a prestação de serviços conforme especificado.",
        "2. DO VALOR: O valor total do contrato será de R$ {{ valor_contrato }}.",
        "3. DO PAGAMENTO: O pagamento será realizado conforme acordado entre as partes."
    ]
    
    for clausula in clausulas:
        doc.add_paragraph(clausula)
    
    # Local e data
    doc.add_paragraph('\nLocal e Data: __________________')
    doc.add_paragraph('\nAssinaturas:')
    doc.add_paragraph('_______________________        _______________________')
    doc.add_paragraph('Contratante                   Contratado')
    
    # Criar diretório de templates se não existir
    os.makedirs('templates_documentos', exist_ok=True)
    
    # Salvar documento
    doc.save('templates_documentos/contrato_servicos_template.docx')
    print("Template de Contrato de Serviços criado com sucesso!")

def criar_template_aluguel():
    doc = Document()
    
    # Título
    titulo = doc.add_heading('CONTRATO DE LOCAÇÃO', 0)
    titulo.alignment = 1  # Centralizado
    
    # Parágrafos
    doc.add_paragraph('Pelo presente instrumento particular de contrato de locação, de um lado,')
    
    # Campos para preenchimento
    doc.add_paragraph('LOCATÁRIO: {{ nome_cliente }}')
    doc.add_paragraph('CPF: {{ cpf }}')
    doc.add_paragraph('LOCADOR: [NOME DO PROPRIETÁRIO]')
    doc.add_paragraph('VALOR DO ALUGUEL: R$ {{ valor_aluguel }}')
    
    # Cláusulas
    doc.add_heading('CLÁUSULAS DO CONTRATO', level=1)
    
    clausulas = [
        "1. DO OBJETO: O presente contrato tem por objeto a locação de imóvel.",
        "2. DO VALOR: O valor mensal do aluguel será de R$ {{ valor_aluguel }}.",
        "3. DO PAGAMENTO: O pagamento será realizado até o 5º dia útil de cada mês."
    ]
    
    for clausula in clausulas:
        doc.add_paragraph(clausula)
    
    # Local e data
    doc.add_paragraph('\nLocal e Data: __________________')
    doc.add_paragraph('\nAssinaturas:')
    doc.add_paragraph('_______________________        _______________________')
    doc.add_paragraph('Locatário                     Locador')
    
    # Salvar documento
    doc.save('templates_documentos/contrato_aluguel_template.docx')
    print("Template de Contrato de Aluguel criado com sucesso!")

if __name__ == '__main__':
    criar_template_contrato_servicos()
    criar_template_aluguel()

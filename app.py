from flask import Flask, render_template, request, send_file, jsonify
from flask_cors import CORS
import os
import sys
import docx
from docx.shared import Pt
from docx2pdf import convert
import uuid

sys.path.append(os.path.dirname(os.path.abspath(__file__)))
from contract_generator import ContractGenerator

app = Flask(__name__)
CORS(app)

contract_generator = ContractGenerator()

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/gerar_contrato', methods=['POST'])
def gerar_contrato():
    dados = request.json
    
    # Validação básica dos dados
    campos_obrigatorios = [
        "nome_cliente", "tipo_evento", "data_evento", 
        "local", "valor_total", "forma_pagamento"
    ]
    
    for campo in campos_obrigatorios:
        if not dados.get(campo):
            return jsonify({"erro": f"Campo {campo} é obrigatório"}), 400
    
    try:
        # Gerar contrato com os dados recebidos
        filename = contract_generator.generate_contract({
            "Nome do Cliente": dados['nome_cliente'],
            "Tipo de Evento": dados['tipo_evento'],
            "Data do Evento": dados['data_evento'],
            "Local": dados['local'],
            "Valor Total": dados['valor_total'],
            "Forma de Pagamento": dados['forma_pagamento']
        })
        
        return jsonify({
            "mensagem": "Contrato gerado com sucesso!", 
            "arquivo": os.path.basename(filename)
        }), 200
    
    except Exception as e:
        return jsonify({"erro": str(e)}), 500

@app.route('/criar_documento', methods=['POST'])
def criar_documento():
    dados = request.json
    
    # Criar documento Word
    doc = docx.Document()
    
    # Estilo do documento
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    # Título
    titulo = doc.add_heading('CONTRATO DE EVENTO', 0)
    titulo.alignment = 1  # Centralizado
    
    # Adicionar informações do contrato
    doc.add_paragraph(f"Nome do Cliente: {dados.get('nome_cliente', 'N/A')}")
    doc.add_paragraph(f"Tipo de Evento: {dados.get('tipo_evento', 'N/A')}")
    doc.add_paragraph(f"Data do Evento: {dados.get('data_evento', 'N/A')}")
    doc.add_paragraph(f"Local: {dados.get('local', 'N/A')}")
    doc.add_paragraph(f"Valor Total: R$ {dados.get('valor_total', 'N/A')}")
    doc.add_paragraph(f"Forma de Pagamento: {dados.get('forma_pagamento', 'N/A')}")
    
    # Adicionar conteúdo do editor
    if dados.get('conteudo_documento'):
        doc.add_paragraph(dados['conteudo_documento'])
    
    # Salvar documento
    nome_arquivo = f"contrato_{uuid.uuid4()}.docx"
    caminho_docx = os.path.join('output', nome_arquivo)
    doc.save(caminho_docx)
    
    # Converter para PDF
    nome_pdf = nome_arquivo.replace('.docx', '.pdf')
    caminho_pdf = os.path.join('output', nome_pdf)
    convert(caminho_docx, caminho_pdf)
    
    return jsonify({
        "mensagem": "Documento criado com sucesso!", 
        "arquivo_docx": nome_arquivo,
        "arquivo_pdf": nome_pdf
    }), 200

@app.route('/listar_contratos')
def listar_contratos():
    contratos = contract_generator.listar_contratos()
    return jsonify(contratos)

@app.route('/download_contrato/<filename>')
def download_contrato(filename):
    return send_file(
        os.path.join('output', filename), 
        as_attachment=True
    )

if __name__ == '__main__':
    # Garantir que a pasta de output exista
    os.makedirs('output', exist_ok=True)
    app.run(host='0.0.0.0', port=5000, debug=True)
